from fpdf import FPDF
from io import BytesIO
from docx import Document
import re

# Accent color for professional look (dark navy)
ACCENT_COLOR = (30, 60, 114)
DARK_GRAY = (50, 50, 50)
MEDIUM_GRAY = (100, 100, 100)
LIGHT_GRAY = (180, 180, 180)

class CV_PDF(FPDF):
    def header(self):
        pass

    def footer(self):
        self.set_y(-15)
        self.set_font("Helvetica", "I", 8)
        self.set_text_color(*MEDIUM_GRAY)
        self.cell(0, 10, f"Page {self.page_no()}/{{nb}}", align="C")

    def _sanitize(self, text):
        """Clean unicode characters for latin-1 encoding."""
        replacements = {
            '\u2022': '-', '\u2019': "'", '\u2018': "'", '\u201c': '"', '\u201d': '"',
            '\u2013': '-', '\u2014': '--', '\u2026': '...', '\u00a0': ' ',
            '\u2014': '--', '\u2012': '-', '\u2010': '-',
        }
        for char, repl in replacements.items():
            text = text.replace(char, repl)
        return text.encode('latin-1', errors='replace').decode('latin-1')

    def _strip_md(self, text):
        """Remove markdown formatting characters."""
        text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
        text = re.sub(r'__(.+?)__', r'\1', text)
        text = re.sub(r'\*(.+?)\*', r'\1', text)
        text = re.sub(r'_(.+?)_', r'\1', text)
        return text

    def _draw_section_rule(self):
        """Draw a horizontal rule under section headers."""
        x = self.get_x()
        y = self.get_y()
        self.set_draw_color(*ACCENT_COLOR)
        self.set_line_width(0.5)
        self.line(x, y, self.w - self.r_margin, y)
        self.ln(4)

    def _write_mixed_line(self, text, size=11, indent=0):
        """Write a line with mixed bold/normal text using ** markers."""
        self.set_x(self.l_margin + indent)
        parts = re.split(r'(\*\*.+?\*\*)', text)
        for part in parts:
            if not part:
                continue
            if part.startswith('**') and part.endswith('**'):
                content = self._sanitize(part[2:-2])
                self.set_font("Helvetica", "B", size)
                self.set_text_color(*DARK_GRAY)
            else:
                content = self._sanitize(part)
                self.set_font("Helvetica", "", size)
                self.set_text_color(*DARK_GRAY)
            # Use write() for inline text, handling page breaks
            if self.get_y() > self.h - self.b_margin - 10:
                self.add_page()
            self.write(6, content)
        self.ln(6)

    def _parse_contact_line(self, line):
        """Parse contact info line with **Label** | Value format."""
        # Remove leading/trailing ** markers
        line = line.strip()
        # Split by |
        parts = [p.strip() for p in line.split('|')]
        return parts

    def _parse_job_header(self, line):
        """Parse **Title** | Company | Location | Dates pattern."""
        # Extract bold title
        title_match = re.match(r'\*\*(.+?)\*\*\s*\|', line)
        if not title_match:
            return None
        title = title_match.group(1)
        rest = line[title_match.end():].strip()
        parts = [p.strip() for p in rest.split('|')]
        return {'title': title, 'parts': parts}

    def print_markdown(self, md_text):
        """Render markdown CV with professional formatting."""
        lines = md_text.split('\n')
        i = 0
        in_contact_block = False
        prev_was_section_header = False

        while i < len(lines):
            line = lines[i]

            # H1: Name
            if line.startswith('# ') and not line.startswith('## '):
                name = self._sanitize(line[2:].strip())
                self.set_font("Helvetica", "B", 20)
                self.set_text_color(*ACCENT_COLOR)
                self.cell(0, 14, name, new_x="LMARGIN", new_y="NEXT")
                self.ln(3)
                prev_was_section_header = True
                i += 1
                continue

            # Contact block: line with | and ** patterns (after name)
            if '|' in line and '**' in line and not line.startswith('#') and not line.startswith('-'):
                contact_parts = self._parse_contact_line(line)
                self.set_font("Helvetica", "", 9)
                self.set_text_color(*MEDIUM_GRAY)
                
                # Build contact string with bold labels
                contact_str = ""
                for idx, part in enumerate(contact_parts):
                    # Check if part has ** markers
                    if '**' in part:
                        # Bold label + value
                        label_match = re.match(r'\*\*(.+?)\*\*\s*(.*)', part)
                        if label_match:
                            label = self._sanitize(label_match.group(1))
                            value = self._sanitize(label_match.group(2).strip())
                            contact_str += f"{label}: {value}"
                        else:
                            contact_str += self._strip_md(part)
                    else:
                        contact_str += self._sanitize(part)
                    
                    if idx < len(contact_parts) - 1:
                        contact_str += "  |  "
                
                # Left-aligned contact line
                self.set_font("Helvetica", "", 9)
                self.set_text_color(*MEDIUM_GRAY)
                self.multi_cell(w=0, h=5, text=contact_str)
                self.ln(2)
                
                # Draw rule under contact block
                self.set_draw_color(*LIGHT_GRAY)
                self.set_line_width(0.3)
                self.line(self.l_margin, self.get_y(), self.w - self.r_margin, self.get_y())
                self.ln(6)
                prev_was_section_header = True
                i += 1
                continue

            # H2: Section headers
            if line.startswith('## '):
                section_title = self._sanitize(line[3:].strip())
                self.ln(4)
                self.set_font("Helvetica", "B", 13)
                self.set_text_color(*ACCENT_COLOR)
                self.cell(0, 9, section_title, new_x="LMARGIN", new_y="NEXT")
                self._draw_section_rule()
                prev_was_section_header = True
                i += 1
                continue

            # Job/Education header: **Title** | Company | Location | Dates
            job_info = self._parse_job_header(line)
            if job_info:
                self.ln(2)
                # Title in bold
                self.set_font("Helvetica", "B", 11)
                self.set_text_color(*DARK_GRAY)
                self.cell(0, 7, self._sanitize(job_info['title']), new_x="LMARGIN", new_y="NEXT")
                
                # Draw divider line below title
                self.set_draw_color(*LIGHT_GRAY)
                self.set_line_width(0.3)
                self.line(self.l_margin, self.get_y(), self.w - self.r_margin, self.get_y())
                self.ln(2)
                
                # Company in bold, location/dates in italic
                if job_info['parts']:
                    company = self._sanitize(job_info['parts'][0])
                    self.set_font("Helvetica", "B", 10)
                    self.set_text_color(*DARK_GRAY)
                    self.cell(0, 6, company, new_x="LMARGIN", new_y="NEXT")
                    
                    if len(job_info['parts']) > 1:
                        meta_parts = [self._sanitize(p) for p in job_info['parts'][1:]]
                        meta_text = "  |  ".join(meta_parts)
                        self.set_font("Helvetica", "I", 9)
                        self.set_text_color(*MEDIUM_GRAY)
                        self.cell(0, 5, meta_text, new_x="LMARGIN", new_y="NEXT")
                self.ln(1)
                prev_was_section_header = False
                i += 1
                continue

            # Bullet points
            if line.startswith('- ') or line.startswith('* '):
                content = line[2:].strip()
                if not content:
                    i += 1
                    continue
                # Check for bold within bullet
                if '**' in content:
                    self.set_x(self.l_margin + 6)
                    self.set_font("Helvetica", "", 10)
                    self.set_text_color(*DARK_GRAY)
                    # Write bullet
                    self.write(5, "-  ")
                    # Write mixed content
                    parts = re.split(r'(\*\*.+?\*\*)', content)
                    for part in parts:
                        if not part:
                            continue
                        if part.startswith('**') and part.endswith('**'):
                            self.set_font("Helvetica", "B", 10)
                            self.write(5, self._sanitize(part[2:-2]))
                        else:
                            self.set_font("Helvetica", "", 10)
                            self.write(5, self._sanitize(part))
                    self.ln(5)
                else:
                    self.set_x(self.l_margin + 6)
                    self.set_font("Helvetica", "", 10)
                    self.set_text_color(*DARK_GRAY)
                    clean = self._sanitize(content)
                    # Use multi_cell for wrapping
                    self.multi_cell(w=self.w - self.l_margin - self.r_margin - 6, h=5, text=f"-  {clean}")
                prev_was_section_header = False
                i += 1
                continue

            # Skills line: **Category:** items
            skills_match = re.match(r'\*\*(.+?):\*\*\s*(.+)', line)
            if skills_match:
                category = self._sanitize(skills_match.group(1))
                items = self._sanitize(skills_match.group(2))
                self.set_font("Helvetica", "B", 10)
                self.set_text_color(*DARK_GRAY)
                self.cell(0, 6, f"{category}:", new_x="LMARGIN", new_y="NEXT")
                self.set_font("Helvetica", "", 10)
                self.set_text_color(*DARK_GRAY)
                self.multi_cell(w=0, h=5, text=f"   {items}")
                self.ln(1)
                prev_was_section_header = False
                i += 1
                continue

            # Empty line: add spacing
            if line.strip() == "":
                if not prev_was_section_header:
                    self.ln(2)
                i += 1
                continue

            # Regular text (summary, descriptions, etc.)
            if line.strip():
                self.set_font("Helvetica", "", 10)
                self.set_text_color(*DARK_GRAY)
                clean = self._sanitize(line)
                self.multi_cell(w=0, h=5, text=clean)
                self.ln(1)
                prev_was_section_header = False

            i += 1


def create_pdf_bytes(text: str) -> bytes:
    pdf = CV_PDF(orientation="P", unit="mm", format="A4")
    pdf.alias_nb_pages()
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.set_margins(left=18, top=18, right=18)
    pdf.add_page()
    pdf.print_markdown(text)
    return bytes(pdf.output())

def create_docx_bytes(text: str) -> bytes:
    doc = Document()
    for line in text.split('\n'):
        if line.startswith('# '):
            doc.add_heading(line[2:], 0)
        elif line.startswith('## '):
            doc.add_heading(line[3:], 1)
        elif line.startswith('### '):
            doc.add_heading(line[4:], 2)
        else:
            p = doc.add_paragraph()
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    p.add_run(part[2:-2]).bold = True
                else:
                    p.add_run(part)
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()
