import streamlit as st
from google import genai
from docx import Document
from fpdf import FPDF
from io import BytesIO
from pypdf import PdfReader
import re
import subprocess
import tempfile
import os
from userCVprofiling import load_profile, get_profile_text, PROFILE_PATH

# --- CONFIG ---
st.set_page_config(page_title="Local AI CV Builder", layout="wide")

# API Setup
st.sidebar.title("Settings")
api_key = st.sidebar.text_input("Gemini API Key", type="password")

# --- Model Discovery ---
available_models = ["gemini-1.5-pro", "gemini-1.5-flash", "gemini-2.0-flash-exp"]
client = None

if api_key:
    try:
        client = genai.Client(api_key=api_key)
        # Fetch models
        models = client.models.list()
        dynamic_models = [m.name for m in models]
        if dynamic_models:
            available_models = dynamic_models
    except Exception as e:
        st.sidebar.error(f"Error fetching models: {e}")

model_name = st.sidebar.selectbox("Select Model", available_models, index=0)

# --- PDF GENERATOR CLASS ---
class CV_PDF(FPDF):
    def header(self):
        pass

    def _sanitize(self, text):
        """Replace Unicode chars that can't be rendered by core fonts."""
        replacements = {
            '\u2022': '-',   # •
            '\u2019': "'",   # '
            '\u2018': "'",   # '
            '\u201c': '"',   # "
            '\u201d': '"',   # "
            '\u2013': '-',   # –
            '\u2014': '--',  # —
            '\u2026': '...', # …
            '\u00a0': ' ',   # non-breaking space
        }
        for char, repl in replacements.items():
            text = text.replace(char, repl)
        # Strip any remaining non-latin-1 characters
        return text.encode('latin-1', errors='replace').decode('latin-1')

    def _break_long_words(self, text, max_chars=80):
        """Break very long unbroken strings to prevent FPDF overflow."""
        words = text.split(' ')
        result = []
        for word in words:
            while len(word) > max_chars:
                result.append(word[:max_chars])
                word = word[max_chars:]
            result.append(word)
        return ' '.join(result)

    def safe_write(self, w, h, txt):
        """Safe wrapper for multi_cell with sanitization and long-word breaking."""
        txt = self._sanitize(txt)
        txt = self._break_long_words(txt)
        try:
            self.multi_cell(w=w, h=h, text=txt, new_x="LMARGIN", new_y="NEXT")
        except Exception:
            for chunk in txt.split('\n'):
                self.cell(w=0, h=h, text=chunk[:100], new_x="LMARGIN", new_y="NEXT")

    def print_markdown(self, md_text):
        self.set_font("Helvetica", size=11)
        lines = md_text.split('\n')
        for line in lines:
            if line.startswith('# '):
                self.set_font("Helvetica", 'B', 16)
                self.cell(w=0, h=10, text=self._sanitize(line[2:]), new_x="LMARGIN", new_y="NEXT")
                self.ln(2)
            elif line.startswith('## '):
                self.set_font("Helvetica", 'B', 14)
                self.cell(w=0, h=10, text=self._sanitize(line[3:]), new_x="LMARGIN", new_y="NEXT")
                self.ln(1)
            elif line.startswith('### '):
                self.set_font("Helvetica", 'B', 12)
                self.cell(w=0, h=8, text=self._sanitize(line[4:]), new_x="LMARGIN", new_y="NEXT")
            elif line.startswith('- ') or line.startswith('* '):
                self.set_font("Helvetica", size=11)
                clean_line = line.replace('**', '').replace('__', '')
                self.safe_write(0, 6, f"  - {clean_line[2:]}")
            else:
                self.set_font("Helvetica", size=11)
                clean_line = line.replace('**', '').replace('__', '')
                self.safe_write(0, 6, clean_line)
                if line.strip() == "": self.ln(2)

def create_pdf_bytes(text):
    pdf = CV_PDF()
    pdf.set_margin(15)
    pdf.add_page()
    pdf.print_markdown(text)
    return bytes(pdf.output())

def create_docx_bytes(text):
    doc = Document()
    for line in text.split('\n'):
        if line.startswith('# '):
            doc.add_heading(line[2:], 0)
        elif line.startswith('## '):
            doc.add_heading(line[3:], 1)
        elif line.startswith('### '):
            doc.add_heading(line[4:], 2)
        else:
            # Simple bold parsing for Word
            p = doc.add_paragraph()
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    p.add_run(part[2:-2]).bold = True
                else:
                    p.add_run(part)
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

# --- APP UI ---
st.title("📄 Local AI CV Tailor")

# ── Saved Profile Detection ──────────────────────────────────────────────────
profile = load_profile()
has_saved_profile = PROFILE_PATH.exists()

if has_saved_profile:
    name = profile.get('personal_info', {}).get('full_name', 'your profile')
    n_jobs = len(profile.get('work_experience', []))
    n_skills = sum(len(profile.get('skills', {}).get(k, [])) for k in ('technical', 'soft', 'languages', 'certifications'))
    st.success(
        f"✅ Saved profile detected — **{name}** "
        f"({n_jobs} roles, {n_skills} skills). "
        f"You can generate a CV without uploading files."
    )

col1, col2 = st.columns([1, 1])

with col1:
    jd = st.text_area("Target Job Description", height=250)
    if has_saved_profile:
        use_profile = st.checkbox("Use saved profile (user_profile.json)", value=True)
    else:
        use_profile = False
    uploads = st.file_uploader(
        "Upload Current CVs (optional if profile is saved)",
        type=['pdf', 'docx', 'doc'],
        accept_multiple_files=True,
    )

with col2:
    if has_saved_profile and use_profile:
        st.info("Using your saved profile. You can still upload extra CVs to supplement it.")
    else:
        st.info("The AI will merge your experience and prioritize skills found in the JD.")
    focus = st.text_input("Special Focus (e.g. 'Focus on Cloud Architecture')")
    generate_btn = st.button("✨ Generate Tailored CV", use_container_width=True)

if generate_btn:
    if not api_key:
        st.error("Please enter API key.")
    elif not jd:
        st.error("Please provide a Job Description.")
    elif not uploads and not use_profile:
        st.error("Please provide at least one CV or enable your saved profile.")
    else:
        with st.spinner("Processing documents..."):
            # Build candidate context from saved profile + any uploads
            context_text = ""

            # 1. Load saved profile text if enabled
            if use_profile and has_saved_profile:
                context_text += get_profile_text(profile)
                context_text += "\n\n"

            # 2. Append text from any uploaded files
            if uploads:
                for f in uploads:
                    if f.name.endswith('.pdf'):
                        reader = PdfReader(f)
                        for page in reader.pages:
                            context_text += page.extract_text()
                    elif f.name.endswith('.doc'):
                        with tempfile.NamedTemporaryFile(suffix='.doc', delete=False) as tmp:
                            tmp.write(f.read())
                            tmp_path = tmp.name
                        try:
                            txt_path = tmp_path.replace('.doc', '.txt')
                            subprocess.run(['textutil', '-convert', 'txt', tmp_path, '-output', txt_path], check=True)
                            with open(txt_path, 'r') as txt_file:
                                context_text += txt_file.read()
                            os.unlink(txt_path)
                        finally:
                            os.unlink(tmp_path)
                    else:
                        doc = Document(f)
                        context_text += "\n".join([p.text for p in doc.paragraphs])
            
            prompt = f"""
            You are an expert CV writer. Create a professional CV in Markdown.
            
            JOB DESCRIPTION:
            {jd}
            
            CANDIDATE DATA:
            {context_text}
            
            REQUIREMENTS:
            - Use Markdown (H1 for Name, H2 for Sections).
            - Focus on: {focus}
            - Match keywords from the JD naturally.
            - Ensure bullet points are achievement-oriented.
            """
            
            response = client.models.generate_content(
                model=model_name,
                contents=prompt
            )
            st.session_state['cv_md'] = response.text

# --- RESULTS & EXPORT ---
if 'cv_md' in st.session_state:
    st.divider()
    editable_md = st.text_area("Review & Edit Markdown", value=st.session_state['cv_md'], height=400)
    
    c1, c2 = st.columns(2)
    with c1:
        st.download_button("Download Word (.docx)", 
                           data=create_docx_bytes(editable_md), 
                           file_name="Tailored_CV.docx", 
                           mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    with c2:
        st.download_button("Download PDF (.pdf)", 
                           data=create_pdf_bytes(editable_md), 
                           file_name="Tailored_CV.pdf", 
                           mime="application/pdf")